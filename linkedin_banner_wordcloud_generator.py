"""LinkedIn Banner Word Cloud Generator.

Reads a resume (PDF, DOCX, or TXT), sends the text to an AI provider to
extract professional terms with importance weights, then renders a word cloud
sized for LinkedIn banners (1584 × 396 px).

Supported AI providers: Claude, ChatGPT, Gemini, Mistral, Groq.
API keys are stored locally at ~/.wordcloud_generator/config.json.
The UI theme follows the OS dark/light mode automatically.

Dependencies (install once)::

    pip install wordcloud matplotlib Pillow pdfplumber python-docx darkdetect

Plus whichever AI provider you want to use::

    pip install anthropic       # Claude
    pip install openai          # ChatGPT
    pip install google-genai    # Gemini
    pip install mistralai       # Mistral
    pip install groq            # Groq
"""
from __future__ import annotations  # lazy type-hint evaluation — Python 3.9+

# ---------------------------------------------------------------------------
# Standard-library imports (all before third-party)
# ---------------------------------------------------------------------------
import json
import logging
import os
import random
import re
import shutil
import string
import sys
import tempfile
import threading
import tkinter as tk
import warnings
from pathlib import Path
from tkinter import filedialog, messagebox, ttk

# ---------------------------------------------------------------------------
# Suppress noisy-but-harmless warnings before any third-party import
# ---------------------------------------------------------------------------
# pdfminer logs FontBBox problems through the *logging* module, not warnings.
logging.getLogger("pdfminer").setLevel(logging.ERROR)
# Pillow / matplotlib can emit DeprecationWarnings about escape sequences.
warnings.filterwarnings(
    "ignore",
    message=r".*invalid escape sequence.*",
    category=DeprecationWarning,
)

# ---------------------------------------------------------------------------
# Application logging
# ---------------------------------------------------------------------------
LOG_DIR = Path.home() / ".wordcloud_generator"
LOG_DIR.mkdir(exist_ok=True)
LOG_FILE = LOG_DIR / "app.log"

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s  %(levelname)-8s  %(message)s",
    datefmt="%Y-%m-%d %H:%M:%S",
    handlers=[
        logging.FileHandler(LOG_FILE, encoding="utf-8"),
        logging.StreamHandler(sys.stdout),
    ],
)
# Keep the console handler quiet — only warnings and errors.
logging.getLogger().handlers[1].setLevel(logging.WARNING)
log = logging.getLogger("wordcloud")
log.info("=" * 60)
log.info("LinkedIn Banner Word Cloud Generator — starting up")
log.info("Python %s on %s", sys.version, sys.platform)


# ---------------------------------------------------------------------------
# Auto-install missing third-party packages
# ---------------------------------------------------------------------------
def _auto_install() -> None:
    """Check for missing packages and install them automatically.

    Tries a plain ``pip install`` first, then falls back to
    ``--break-system-packages`` (needed on Python 3.12+ Debian/Ubuntu/Fedora)
    and finally ``--user`` as a last resort.
    """
    import subprocess  # imported here to avoid polluting the module namespace

    deps: dict[str, str] = {
        "wordcloud":  "wordcloud",
        "matplotlib": "matplotlib",
        "PIL":        "Pillow",
        "pdfplumber": "pdfplumber",
        "docx":       "python-docx",
        "darkdetect": "darkdetect",
    }
    missing = [
        pip_name
        for mod, pip_name in deps.items()
        if not _can_import(mod)
    ]
    if not missing:
        return

    print(f"\n  Installing {len(missing)} missing package(s): {', '.join(missing)}")
    print("  (one-time setup — may take a moment)\n")

    base = [sys.executable, "-m", "pip", "install", "--quiet"]
    for extra in [[], ["--break-system-packages"], ["--user"]]:
        try:
            subprocess.check_call(base + extra + missing)
            print()
            return
        except subprocess.CalledProcessError:
            continue

    print(
        f"\n  Auto-install failed. Please run manually:\n"
        f"  pip install {' '.join(missing)}\n"
    )
    sys.exit(1)


def _can_import(module_name: str) -> bool:
    """Return True if *module_name* can be imported without error."""
    try:
        __import__(module_name)
        return True
    except ImportError:
        return False


_auto_install()

# Third-party imports — must come after _auto_install().
import matplotlib                   # noqa: E402
matplotlib.use("Agg")               # non-interactive backend, safe for threads
import matplotlib.pyplot as plt     # noqa: E402
from PIL import Image, ImageTk      # noqa: E402
from wordcloud import WordCloud     # noqa: E402

from providers import PROVIDERS, get_provider  # noqa: E402


# ---------------------------------------------------------------------------
# OS dark-mode detection
# ---------------------------------------------------------------------------
def _detect_dark_mode() -> bool:
    """Return ``True`` if the OS is running in dark mode.

    Tries ``darkdetect`` first (cross-platform), then falls back to
    platform-specific queries for macOS and Windows.
    """
    try:
        import darkdetect
        return bool(darkdetect.isDark())
    except Exception:
        pass

    if sys.platform == "darwin":
        try:
            import subprocess
            result = subprocess.run(
                ["defaults", "read", "-g", "AppleInterfaceStyle"],
                capture_output=True,
                text=True,
            )
            return result.stdout.strip().lower() == "dark"
        except Exception:
            pass

    if sys.platform == "win32":
        try:
            import winreg
            key = winreg.OpenKey(
                winreg.HKEY_CURRENT_USER,
                r"Software\Microsoft\Windows\CurrentVersion\Themes\Personalize",
            )
            val, _ = winreg.QueryValueEx(key, "AppsUseLightTheme")
            return val == 0
        except Exception:
            pass

    return False


_DARK = _detect_dark_mode()
log.info("Theme: %s", "dark" if _DARK else "light")

# ---------------------------------------------------------------------------
# Colour themes
# ---------------------------------------------------------------------------
_THEMES: dict[str, dict[str, str]] = {
    "light": dict(
        bg           = "#e8ecf1",
        frame        = "#ffffff",
        primary      = "#1d4ed8",
        primary_hv   = "#1e40af",
        danger       = "#b91c1c",
        danger_hv    = "#991b1b",
        success      = "#15803d",
        success_hv   = "#166534",
        text         = "#0f172a",
        text_light   = "#475569",
        tree_bg      = "#ffffff",
        tree_fg      = "#0f172a",
        tree_sel     = "#bfdbfe",
        tree_head_bg = "#cbd5e1",
        tree_head_fg = "#0f172a",
        entry_bg     = "#ffffff",
        border       = "#cbd5e1",
    ),
    "dark": dict(
        bg           = "#0f172a",
        frame        = "#1e293b",
        primary      = "#3b82f6",
        primary_hv   = "#2563eb",
        danger       = "#ef4444",
        danger_hv    = "#dc2626",
        success      = "#22c55e",
        success_hv   = "#16a34a",
        text         = "#f1f5f9",
        text_light   = "#94a3b8",
        tree_bg      = "#1e293b",
        tree_fg      = "#f1f5f9",
        tree_sel     = "#1e3a5f",
        tree_head_bg = "#0f172a",
        tree_head_fg = "#94a3b8",
        entry_bg     = "#1e293b",
        border       = "#334155",
    ),
}

#: Active colour palette — every widget references ``C["key"]``.
C: dict[str, str] = _THEMES["dark" if _DARK else "light"]


# ---------------------------------------------------------------------------
# Application constants
# ---------------------------------------------------------------------------
BANNER_WIDTH: int   = 1584       # LinkedIn banner width in pixels
BANNER_HEIGHT: int  = 396        # LinkedIn banner height in pixels
MAX_FILE_SIZE_MB: float = 3.0    # PNG size cap; exceeded → colour-quantise
MAX_TERMS: int      = 60         # Maximum terms shown in the treeview
WEIGHT_MIN: int     = 1_000      # Minimum term weight passed to the AI
WEIGHT_MAX: int     = 10_000     # Maximum term weight passed to the AI
MANUAL_ADD_WEIGHT: int = 7_000   # Default weight for user-added terms
MAX_RESUME_CHARS: int = 15_000   # Resume text truncation limit for API calls

WINDOW_MIN_WIDTH: int  = 1060
WINDOW_MIN_HEIGHT: int = 720


# ---------------------------------------------------------------------------
# Config helpers
# ---------------------------------------------------------------------------
CONFIG_FILE: Path = LOG_DIR / "config.json"


def load_config() -> dict:
    """Load the application config from disk.

    Returns a default structure if the file does not exist or is corrupt.
    """
    if not CONFIG_FILE.exists():
        return {"active_provider": "", "keys": {}}
    try:
        return json.loads(CONFIG_FILE.read_text(encoding="utf-8"))
    except (json.JSONDecodeError, OSError):
        return {"active_provider": "", "keys": {}}


def save_config(config: dict) -> None:
    """Persist *config* to disk as pretty-printed JSON."""
    CONFIG_FILE.write_text(json.dumps(config, indent=2), encoding="utf-8")


# ---------------------------------------------------------------------------
# AI term-extraction prompt
# ---------------------------------------------------------------------------
TERM_EXTRACTION_PROMPT: str = """\
You are an expert resume analyst with deep knowledge across all professional \
fields — technology, finance, healthcare, law, education, marketing, \
engineering, science, creative industries, government, non-profit, and more.

Your task: read the resume below, identify the person's industry and role, \
then extract the terms that best represent their professional identity for a \
LinkedIn banner word cloud.

Return ONLY a JSON object — no markdown fences, no preamble, no explanation. \
Your entire response must be valid JSON and nothing else.

{{
  "terms": [
    {{"term": "Term One", "weight": 9500}},
    {{"term": "Term Two", "weight": 7800}},
    ...up to 60 entries...
  ],
  "runner_ups": [
    {{"term": "Term Three", "weight": 4200}},
    ...up to 30 entries...
  ]
}}

━━━  WHAT TO EXTRACT  ━━━

Extract whichever of the following are relevant to THIS person's field. \
Do not force tech terms onto a non-tech resume, or medical terms onto a \
marketing resume. Let the resume guide you.

• Hard skills — tools, software, platforms, instruments, techniques, methods \
  (examples across industries: Excel, Photoshop, AutoCAD, R, QuickBooks, \
  Salesforce, Pro Tools, SPSS, LexisNexis, Epic, MATLAB, Final Cut Pro)
• Domain expertise — areas of specialization, practice areas, subject matter \
  (examples: Contract Law, Oncology, Brand Strategy, Structural Engineering, \
  Financial Modeling, Curriculum Design, Clinical Trials, Supply Chain)
• Methodologies and frameworks — how the person works \
  (examples: Agile, Six Sigma, Design Thinking, Lean, Scrum, \
  Case Management, Evidence-Based Practice, Montessori, Value Investing)
• Industry credentials and standards — certifications, protocols, regulations \
  (examples: CPA, PMP, SHRM, Bar Admission, HIPAA, GAAP, ISO 9001, \
  Series 7, LEED, AWS Certified, Google Analytics, Cisco CCNA)
• Professional concepts — ideas and practices central to their work \
  (examples: Patient Advocacy, Risk Assessment, Stakeholder Engagement, \
  Fiscal Policy, Grant Writing, UX Research, Conflict Resolution)

━━━  WEIGHTING GUIDE  ━━━

Weight reflects how central a skill is to the person's professional identity \
— not just how often it appears.

  9000–10000  Signature skills: what this person IS known for. The things \
              that would appear in their headline or elevator pitch. \
              These should be few (3–8 terms).
  7000–8999   Core competencies: used regularly, clearly demonstrated with \
              results, central to their day-to-day role.
  5000–6999   Strong supporting skills: mentioned with specifics, part of \
              their toolkit but not the headline.
  3000–4999   Relevant but secondary: present and genuine, adds breadth.
  1000–2999   Peripheral: briefly mentioned, supporting context.

Calibrate weights relative to each other. If this is a software engineer, \
their primary language might be 9500. If this is a surgeon, their specialty \
might be 9800. If this is a marketer, their channel expertise might be 9200.

━━━  COMPOUND TERMS  ━━━

- Each term is 1 or 2 words maximum.
- Prefer specific compound terms over vague single words: \
  "Financial Modeling" beats "Finance", "Patient Care" beats "Care", \
  "Machine Learning" beats "Learning", "Contract Law" beats "Law".
- If you include a compound term, do NOT also list its bare components \
  as separate entries. "Power BI" covers "Power" and "BI" — do not list \
  them separately. Only list a component alone if it carries independent \
  meaning beyond any compound (e.g. "Python" is valid alongside "Python GUI").

━━━  CASING  ━━━

Use each term's real-world official casing — not blind title-case:
• All-caps for initialisms and acronyms: SQL, API, CPA, MBA, HIPAA, \
  GAAP, UX, AI, ML, ROI, KPI, EHR, CAD, SEO, PR, HR, CFO, NGO, etc.
• Brand/product casing: PowerPoint, QuickBooks, LinkedIn, SharePoint, \
  macOS, iOS, GitHub, ChatGPT, Salesforce, HubSpot, WordPress, etc.
• Lowercase brands that style themselves lowercase: pandas, numpy, npm, \
  dbt, pytest, etc.
• Title Case for multi-word domain terms and professional concepts: \
  Financial Modeling, Patient Care, Brand Strategy, Risk Management, etc.

━━━  EXCLUDE  ━━━

• Personal info: name, city, state, country, email, phone, URLs, \
  social handles
• Section headings: Education, Experience, Skills, Summary, References, etc.
• Action verbs: Managed, Led, Developed, Created, Improved, Responsible, etc.
• Degree terms: Bachelor, Master, PhD, MBA, Associate, Diploma, etc.
• Dates, years, month names, time ranges
• Generic filler: Various, Multiple, Several, Other, General, Extensive, etc.
• Unsupported soft skills: "Team Player", "Hard Worker", "Detail Oriented" \
  — unless the resume demonstrates them through specific achievements

━━━  runner_ups  ━━━

"runner_ups" are genuine skills from the resume that didn't make the \
top 60 — not invented terms. They must appear in the actual document.

Resume:
{resume_text}
"""


# ---------------------------------------------------------------------------
# NLP / AI helpers
# ---------------------------------------------------------------------------
def _detect_lang_simple(text: str) -> str:
    """Return a BCP-47-style language code based on script frequency.

    Only used for selecting a system font for the word cloud — full
    language detection is handled by the AI prompt itself.

    Returns ``"ja"``, ``"ko"``, ``"zh-cn"``, or ``"en"`` (default).
    """
    if not text:
        return "en"
    chars = len(text)
    cjk    = sum(1 for c in text if "\u4e00" <= c <= "\u9fff")
    kana   = sum(1 for c in text if "\u3040" <= c <= "\u30ff")
    hangul = sum(1 for c in text if "\uac00" <= c <= "\ud7af")
    if kana   / chars > 0.03:
        return "ja"
    if hangul / chars > 0.03:
        return "ko"
    if cjk    / chars > 0.05:
        return "zh-cn"
    return "en"


def extract_terms_ai(
    resume_text: str,
    provider,
    separator: str = "_",
    progress_callback=None,
) -> tuple[dict[str, int], dict[str, int], str, str]:
    """Send *resume_text* to *provider* and return weighted professional terms.

    Args:
        resume_text:       Raw text extracted from the resume file.
        provider:          An initialised ``BaseProvider`` instance.
        separator:         ``"_"`` or ``" "`` — joined between words in
                           multi-word terms for display.
        progress_callback: Optional ``callable(str)`` for status updates.

    Returns:
        A 4-tuple of ``(terms, runner_ups, lang, provider_name)`` where
        *terms* and *runner_ups* are ``{term: weight}`` dicts.
    """
    def _ui(msg: str) -> None:
        if progress_callback:
            progress_callback(msg)

    truncated = resume_text[:MAX_RESUME_CHARS]
    if len(resume_text) > MAX_RESUME_CHARS:
        log.info(
            "Resume truncated from %d to %d chars for API call",
            len(resume_text), MAX_RESUME_CHARS,
        )

    prompt = TERM_EXTRACTION_PROMPT.format(resume_text=truncated)
    _ui(f"Sending to {provider.name} \u2026")
    log.info(
        "extract_terms_ai: %d chars, provider=%s, model=%s",
        len(truncated), provider.name, provider.model,
    )

    raw = provider.explain(prompt)
    log.info("AI response: %d chars", len(raw))
    _ui("Parsing response \u2026")

    # Strip markdown fences if the model wrapped its JSON response anyway.
    clean = raw.strip()
    if clean.startswith("```"):
        clean = re.sub(r"^```[a-zA-Z]*\n?", "", clean)
        clean = re.sub(r"\n?```\s*$", "", clean).strip()

    try:
        data = json.loads(clean)
    except json.JSONDecodeError:
        # Last-ditch: try to find the outermost {...} block.
        match = re.search(r"\{.*\}", clean, re.DOTALL)
        if match:
            try:
                data = json.loads(match.group())
            except json.JSONDecodeError:
                raise ValueError(
                    "The AI returned a response that could not be parsed as "
                    "JSON. Please try again."
                )
        else:
            raise ValueError("The AI did not return JSON. Please try again.")

    def _parse_entries(
        key: str,
        default_weight: int,
        exclude: dict[str, int] | None = None,
    ) -> dict[str, int]:
        """Parse a list of {term, weight} entries from *data[key]*."""
        result: dict[str, int] = {}
        for entry in data.get(key, []):
            raw_term = str(entry.get("term", "")).strip()
            weight = int(entry.get("weight", default_weight))
            if not raw_term:
                continue
            cleaned = clean_term(raw_term, separator=separator)
            if cleaned and (exclude is None or cleaned not in exclude):
                result[cleaned] = max(WEIGHT_MIN, min(WEIGHT_MAX, weight))
        return result

    terms     = _parse_entries("terms",     default_weight=5000)
    runner_ups = _parse_entries("runner_ups", default_weight=3000, exclude=terms)

    if not terms:
        raise ValueError(
            "The AI returned no terms. "
            "Check that the file is a resume and try again."
        )

    # Remove single-word terms already covered by a higher-weight compound,
    # e.g. standalone "BI" when "Power BI" is already in the list.
    terms      = _dedupe_subword_terms(terms)
    runner_ups = _dedupe_subword_terms(runner_ups, reference=terms)

    lang = _detect_lang_simple(resume_text)
    log.info(
        "Parsed: %d terms, %d runner-ups, lang=%s",
        len(terms), len(runner_ups), lang,
    )
    _ui(f"Done \u2014 {len(terms)} terms, {len(runner_ups)} runner-ups.")
    return terms, runner_ups, lang, provider.name


def _dedupe_subword_terms(
    terms: dict[str, int],
    reference: dict[str, int] | None = None,
) -> dict[str, int]:
    """Remove single-word terms already represented by a higher-weight compound.

    A single-word term *T* is dropped when there exists a compound term
    (2 words) that:

    - contains *T* as one of its component words (case-insensitive), **and**
    - has weight >= *T*'s weight (i.e. the compound is at least as important).

    The ``reference`` pool is checked alongside ``terms`` when deduplicating
    runner-ups against already-accepted main terms.

    Examples::

        "BI" (7400)  → dropped   ("Power BI" at 9800 covers it)
        "SQL" (8800) → kept      (no compound containing SQL outweighs it)
        "Python" (8700) → kept   ("Python GUI" at 2000 does not outweigh it)
    """
    pool = dict(terms)
    if reference:
        pool.update(reference)

    result: dict[str, int] = {}
    for term, weight in terms.items():
        parts = re.split(r"[ _]+", term)
        if len(parts) != 1:
            # Compound terms are always kept.
            result[term] = weight
            continue

        t_lower = term.lower()
        covered = any(
            other != term
            and len(re.split(r"[ _]+", other)) > 1
            and t_lower in [p.lower() for p in re.split(r"[ _]+", other)]
            and other_w >= weight
            for other, other_w in pool.items()
        )
        if not covered:
            result[term] = weight
        else:
            log.debug("Deduped subword '%s' (covered by a compound)", term)

    return result


def _validate_content(text: str) -> tuple[bool, str]:
    """Check that extracted text is plausible resume content.

    Returns a ``(ok, reason)`` tuple. When ``ok`` is ``False``, ``reason``
    is a user-friendly explanation of the problem.
    """
    stripped = text.strip()
    n = len(stripped)

    if n == 0:
        return False, (
            "No text could be extracted from this file.\n\n"
            "If this is a scanned PDF (an image of a document rather than a "
            "text-based PDF), the AI will have nothing to read.\n\n"
            "Try exporting your resume as a text-based PDF from Word, "
            "Google Docs, or Acrobat."
        )

    if n < 150:
        return False, (
            f"Only {n} characters were extracted — unusually short for a "
            f"resume.\n\n"
            f"The file may be image-based, password-protected, or not a resume."
        )

    # High proportion of non-printable / replacement characters indicates
    # binary data that slipped through the encoding fallback.
    non_print = sum(1 for c in stripped if ord(c) < 32 and c not in "\n\r\t")
    replacements = stripped.count("\ufffd")  # Unicode replacement character
    garbage_ratio = (non_print + replacements) / n
    if garbage_ratio > 0.05:
        return False, (
            f"The extracted text contains {garbage_ratio:.0%} unreadable "
            f"characters.\n\n"
            f"The file may be password-protected, corrupted, or in a format "
            f"that cannot be read as plain text."
        )

    return True, ""


# ---------------------------------------------------------------------------
# File reading
# ---------------------------------------------------------------------------
def extract_text_from_file(file_path: str) -> str:
    """Read raw text from a PDF, DOCX, or TXT file.

    Args:
        file_path: Absolute or relative path to the source file.

    Returns:
        The extracted plain-text content.

    Raises:
        RuntimeError: If the file extension is not supported.
    """
    ext = os.path.splitext(file_path)[1].lower()
    log.info("Reading file: %s (type: %s)", file_path, ext)

    if ext == ".pdf":
        import pdfplumber
        with pdfplumber.open(file_path) as pdf:
            text = "\n".join(p.extract_text() or "" for p in pdf.pages)
        log.info("PDF: %d pages, %d chars", len(pdf.pages), len(text))
        return text

    if ext == ".docx":
        import docx
        doc = docx.Document(file_path)
        text = "\n".join(p.text for p in doc.paragraphs)
        log.info("DOCX: %d paragraphs, %d chars", len(doc.paragraphs), len(text))
        return text

    if ext == ".txt":
        with open(file_path, "r", encoding="utf-8") as fh:
            text = fh.read()
        log.info("TXT: %d chars", len(text))
        return text

    raise RuntimeError(f"Unsupported file type: {ext}")


# ---------------------------------------------------------------------------
# Term formatting
# ---------------------------------------------------------------------------
def clean_term(term: str, separator: str = "_") -> str:
    """Normalise a term string for display.

    Strips surrounding punctuation, then rejoins words with *separator*.
    Casing is intentionally preserved exactly as the AI returns it — the
    prompt instructs the AI to use official brand casing.

    Args:
        term:      Raw term string from the AI response.
        separator: ``"_"`` or ``" "`` to join multi-word terms.

    Returns:
        The cleaned term, or an empty string if nothing remains.
    """
    cjk_punct = "\u3002\u3001\uff0c\uff1b\uff1a\uff01\uff1f"
    term = term.strip().strip(string.punctuation + cjk_punct).strip()
    if not term:
        return ""
    parts = [p for p in re.split(r"[\s_]+", term) if p]
    return separator.join(parts)


# ---------------------------------------------------------------------------
# Word cloud generation
# ---------------------------------------------------------------------------
def find_font(lang: str) -> str | None:
    """Return a system font path suitable for *lang*, or ``None`` for Latin.

    Checks standard installation paths on Linux, macOS, and Windows for
    CJK, Arabic, and Devanagari fonts.
    """
    font_map: dict[str, list[str]] = {
        "cjk": [
            "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",
            "/System/Library/Fonts/PingFang.ttc",
            "C:\\Windows\\Fonts\\msyh.ttc",
        ],
        "arabic": [
            "/usr/share/fonts/truetype/noto/NotoSansArabic-Regular.ttf",
            "C:\\Windows\\Fonts\\arial.ttf",
        ],
        "devanagari": [
            "/usr/share/fonts/truetype/noto/NotoSansDevanagari-Regular.ttf",
            "C:\\Windows\\Fonts\\mangal.ttf",
        ],
    }

    if lang in ("zh-cn", "zh-tw", "zh", "ja", "ko"):
        family = "cjk"
    elif lang in ("ar", "fa", "ur", "he"):
        family = "arabic"
    elif lang in ("hi", "mr", "ne"):
        family = "devanagari"
    else:
        return None

    for path in font_map[family]:
        if os.path.isfile(path):
            return path
    return None


def generate_wordcloud(
    terms: dict[str, int],
    background_color: str = "white",
    colormap: str = "turbo",
    output_path: str = "linkedin_banner.png",
    lang: str = "en",
) -> str:
    """Render a LinkedIn-banner word cloud and save it as a PNG.

    Applies a ±10% weight jitter for visual variety, quantises the output
    if it exceeds ``MAX_FILE_SIZE_MB``, and returns the final output path.

    Args:
        terms:            ``{display_term: weight}`` dict.
        background_color: ``"white"`` or ``"black"``.
        colormap:         Any matplotlib colormap name.
        output_path:      Destination file path for the PNG.
        lang:             BCP-47 language code used to select a system font.

    Returns:
        The path to the saved PNG file.
    """
    log.info(
        "generate_wordcloud: %d terms, bg=%s, cmap=%s, lang=%s",
        len(terms), background_color, colormap, lang,
    )

    # Jitter weights by ±10% to vary word sizes slightly across generations.
    jittered = {
        term: random.randint(
            max(100, w - int(w * 0.10)),
            min(WEIGHT_MAX, w + int(w * 0.10)),
        )
        for term, w in terms.items()
    }

    wc_kwargs: dict = dict(
        width=BANNER_WIDTH,
        height=BANNER_HEIGHT,
        background_color=background_color,
        colormap=colormap,
        prefer_horizontal=0.7,
        min_font_size=8,
        max_words=200,
        collocations=False,
        normalize_plurals=False,
        regexp=r"[\w\-\+\#']+",
    )
    font_path = find_font(lang)
    if font_path:
        wc_kwargs["font_path"] = font_path

    wc = WordCloud(**wc_kwargs).generate_from_frequencies(jittered)
    fig, ax = plt.subplots(figsize=(15.84, 3.96))
    ax.imshow(wc, interpolation="bilinear")
    ax.axis("off")
    plt.tight_layout(pad=0)
    plt.savefig(output_path, format="png", dpi=300,
                bbox_inches="tight", pad_inches=0)
    plt.close(fig)

    # Quantise to 256 colours if the lossless PNG is above the size cap.
    file_mb = os.path.getsize(output_path) / (1024 * 1024)
    if file_mb > MAX_FILE_SIZE_MB:
        log.info("PNG too large (%.2f MB), quantising to 256 colours", file_mb)
        img = Image.open(output_path)
        img = img.quantize(colors=256, method=Image.Quantize.MEDIANCUT)
        img.save(output_path, format="png", optimize=True)
        file_mb = os.path.getsize(output_path) / (1024 * 1024)

    log.info("Saved: %s (%.2f MB)", output_path, file_mb)
    return output_path


# ============================================================================
# GUI
# ============================================================================

#: Platform-aware font family for consistent rendering on macOS/Win/Linux.
_FONT_FAMILY: str = "Segoe UI"
if sys.platform == "darwin":
    _FONT_FAMILY = "Helvetica Neue"
elif sys.platform.startswith("linux"):
    _FONT_FAMILY = "DejaVu Sans"


def _downloads_folder() -> Path:
    """Return the best available save destination.

    Priority: ``/data`` (Docker volume) → ``~/Downloads`` → ``~``.
    """
    docker = Path("/data")
    if docker.is_dir():
        return docker
    downloads = Path.home() / "Downloads"
    if downloads.is_dir():
        return downloads
    return Path.home()


def _shorten_path(path_str: str, max_len: int = 55) -> str:
    """Return the filename portion of *path_str*, truncated if necessary.

    Appends a horizontal ellipsis (…) when the name exceeds *max_len*.
    """
    name = Path(path_str).name
    return name if len(name) <= max_len else name[: max_len - 1] + "\u2026"


def _friendly_api_error(exc: Exception) -> str:
    """Translate a raw API exception into a concise user-facing message."""
    msg = str(exc).lower()
    if "credit" in msg or "quota" in msg or "insufficient" in msg:
        return "No credits \u2014 add credits at the provider\u2019s billing page."
    if "invalid" in msg and "key" in msg:
        return "Invalid key \u2014 double-check you copied the whole thing."
    if "rate limit" in msg:
        return "Rate limit hit \u2014 wait a moment and try again."
    if "401" in msg:
        return "Authentication failed \u2014 key may be wrong or expired."
    if "403" in msg:
        return "Access denied \u2014 check your account has API access enabled."
    if "timeout" in msg or "timed out" in msg:
        return "Connection timed out \u2014 check your internet and try again."
    raw = str(exc)
    return raw[:120] + "\u2026" if len(raw) > 120 else raw


# ---------------------------------------------------------------------------
# Settings dialog
# ---------------------------------------------------------------------------
class SettingsDialog(tk.Toplevel):
    """Two-panel dialog for selecting an AI provider and entering an API key.

    Displays a provider picker, masked key entry with a Show/Hide toggle,
    a live connection-test button, and a Save button that writes the key to
    ``config.json`` and updates the main window's provider label.
    """

    #: Per-provider key hints and taglines shown in the dialog.
    _HINTS: dict[str, tuple[str, str]] = {
        "Claude":  (
            "Starts with  sk-ant-api03-\u2026",
            "Anthropic \u00b7 best for nuanced extraction",
        ),
        "ChatGPT": (
            "Starts with  sk-proj-\u2026  or  sk-\u2026",
            "OpenAI \u00b7 widely used",
        ),
        "Gemini":  (
            "Starts with  AIza\u2026",
            "Google \u00b7 free tier available",
        ),
        "Mistral": (
            "A long random string of letters and numbers",
            "Mistral AI \u00b7 efficient European model",
        ),
        "Groq":    (
            "Starts with  gsk_\u2026",
            "Groq \u00b7 extremely fast, generous free tier",
        ),
    }

    def __init__(self, parent: WordCloudApp) -> None:
        super().__init__(parent)
        self.parent_app = parent
        self.title("Settings \u2014 AI Provider")
        self.configure(bg=C["bg"])
        self.resizable(False, False)
        self.grab_set()  # modal

        _f  = _FONT_FAMILY
        _fn = (_f, 10)
        _fb = (_f, 10, "bold")
        _fs = (_f, 9)
        style = ttk.Style(self)

        # ── Header ───────────────────────────────────────────────────────
        tk.Label(
            self, text="AI Provider Settings",
            bg=C["bg"], fg=C["text"], font=(_f, 14, "bold"),
        ).pack(anchor="w", padx=24, pady=(20, 2))
        tk.Label(
            self, text="Select a provider and paste your API key.",
            bg=C["bg"], fg=C["text_light"], font=_fs,
        ).pack(anchor="w", padx=24, pady=(0, 12))

        # ── Provider picker ───────────────────────────────────────────────
        pf = tk.Frame(
            self, bg=C["frame"],
            highlightbackground=C["border"], highlightthickness=1,
        )
        pf.pack(fill="x", padx=24, pady=(0, 12))

        tk.Label(
            pf, text="Provider:", bg=C["frame"], fg=C["text"], font=_fb,
        ).grid(row=0, column=0, sticky="w", padx=12, pady=10)

        config = load_config()
        initial = config.get("active_provider", "") or list(PROVIDERS.keys())[0]
        self._provider_var = tk.StringVar(value=initial)
        self._provider_combo = ttk.Combobox(
            pf,
            values=list(PROVIDERS.keys()),
            textvariable=self._provider_var,
            state="readonly",
            width=18,
        )
        self._provider_combo.grid(row=0, column=1, sticky="w",
                                  padx=(0, 12), pady=10)
        self._provider_combo.bind("<<ComboboxSelected>>",
                                  self._on_provider_change)

        self._tagline_var = tk.StringVar()
        tk.Label(
            pf, textvariable=self._tagline_var,
            bg=C["frame"], fg=C["text_light"], font=_fs,
        ).grid(row=0, column=2, sticky="w", padx=(0, 12), pady=10)

        # ── Key entry ─────────────────────────────────────────────────────
        kf = tk.Frame(
            self, bg=C["frame"],
            highlightbackground=C["border"], highlightthickness=1,
        )
        kf.pack(fill="x", padx=24, pady=(0, 8))

        tk.Label(
            kf, text="API Key:", bg=C["frame"], fg=C["text"], font=_fb,
        ).grid(row=0, column=0, sticky="w", padx=12, pady=(12, 4))

        self._key_var  = tk.StringVar()
        self._show_key = tk.BooleanVar(value=False)
        self._key_entry = tk.Entry(
            kf,
            textvariable=self._key_var,
            width=44,
            show="\u2022",
            font=_fn,
            bg=C["entry_bg"],
            fg=C["text"],
            insertbackground=C["text"],
            relief="flat",
            highlightbackground=C["border"],
            highlightthickness=1,
        )
        self._key_entry.grid(row=0, column=1, sticky="w",
                             padx=(0, 8), pady=(12, 4))

        style.configure("Small.TButton", font=(_f, 8), padding=(6, 3))
        ttk.Button(
            kf, text="Show", style="Small.TButton",
            command=self._toggle_show,
        ).grid(row=0, column=2, padx=(0, 12), pady=(12, 4))

        self._hint_var = tk.StringVar()
        tk.Label(
            kf, textvariable=self._hint_var,
            bg=C["frame"], fg=C["text_light"], font=_fs,
        ).grid(row=1, column=1, columnspan=2,
               sticky="w", padx=(0, 12), pady=(0, 12))

        # ── Status line ───────────────────────────────────────────────────
        self._status_var = tk.StringVar()
        tk.Label(
            self, textvariable=self._status_var,
            bg=C["bg"], fg=C["text_light"], font=(_f, 9, "italic"),
        ).pack(anchor="w", padx=24, pady=(0, 4))

        # ── Buttons ───────────────────────────────────────────────────────
        bf = tk.Frame(self, bg=C["bg"])
        bf.pack(fill="x", padx=24, pady=(4, 20))

        for btn_name, bg_, hv_, fg_ in [
            ("Primary.TButton", C["primary"], C["primary_hv"], "#ffffff"),
            ("Success.TButton", C["success"], C["success_hv"], "#ffffff"),
        ]:
            style.configure(
                btn_name, background=bg_, foreground=fg_,
                font=(_f, 9, "bold"), padding=(14, 5),
                relief="flat", borderwidth=0,
            )
            style.map(
                btn_name,
                background=[("active", hv_), ("pressed", hv_)],
                foreground=[("active", fg_), ("pressed", fg_)],
            )

        ttk.Button(
            bf, text="Test Connection", style="Primary.TButton",
            command=self._test,
        ).pack(side="left", padx=(0, 8))
        ttk.Button(
            bf, text="Save", style="Success.TButton",
            command=self._save,
        ).pack(side="left")
        ttk.Button(bf, text="Cancel", command=self.destroy).pack(side="right")

        # Centre the dialog over its parent window.
        self._on_provider_change()
        self.update_idletasks()
        w = self.winfo_reqwidth()
        h = self.winfo_reqheight()
        x = parent.winfo_x() + (parent.winfo_width()  - w) // 2
        y = parent.winfo_y() + (parent.winfo_height() - h) // 2
        self.geometry(f"{w}x{h}+{x}+{y}")

    # ── Event handlers ────────────────────────────────────────────────────

    def _on_provider_change(self, *_) -> None:
        """Populate the key field and hints when the provider selection changes."""
        name = self._provider_var.get()
        config = load_config()
        self._key_var.set(config.get("keys", {}).get(name, ""))
        hint, tagline = self._HINTS.get(name, ("", ""))
        self._hint_var.set(hint)
        self._tagline_var.set(tagline)
        self._status_var.set("")

    def _toggle_show(self) -> None:
        """Toggle the API key entry between masked and plain-text display."""
        self._show_key.set(not self._show_key.get())
        self._key_entry.config(show="" if self._show_key.get() else "\u2022")

    def _test(self) -> None:
        """Fire a minimal API call to verify the entered key, in a thread."""
        name = self._provider_var.get()
        key  = self._key_var.get().strip()
        if not key:
            self._status_var.set("Paste an API key first.")
            return
        self._status_var.set("Testing connection \u2026")
        self.update_idletasks()

        def _worker() -> None:
            try:
                ok = get_provider(name, key).test_connection()
                msg = (
                    f"\u2713  Connected to {name}." if ok
                    else f"\u2717  {name} returned an empty response."
                )
            except Exception as exc:
                msg = f"\u2717  {_friendly_api_error(exc)}"
            self.after(0, self._status_var.set, msg)

        threading.Thread(target=_worker, daemon=True).start()

    def _save(self) -> None:
        """Save the current provider and key to config, then close."""
        name = self._provider_var.get()
        key  = self._key_var.get().strip()
        if not key:
            self._status_var.set("Enter an API key before saving.")
            return
        config = load_config()
        config.setdefault("keys", {})[name] = key
        config["active_provider"] = name
        save_config(config)
        log.info("Saved API key for provider: %s", name)
        self.parent_app._refresh_provider_label()
        self.destroy()


# ---------------------------------------------------------------------------
# Main application window
# ---------------------------------------------------------------------------
class WordCloudApp(tk.Tk):
    """Top-level application window for the LinkedIn Banner Word Cloud Generator.

    Layout (bottom-anchored):
        - Preview image panel (always visible)
        - Generate / Save / Export action buttons (always visible)
        - Appearance section (hidden until terms are extracted)
        - Scrollable upper area: file picker, term treeview, edit controls
    """

    #: File extensions accepted by Browse and Analyse Resume.
    _SUPPORTED_EXTENSIONS: frozenset[str] = frozenset({".pdf", ".docx", ".txt"})

    def __init__(self) -> None:
        super().__init__()
        self.title("LinkedIn Banner Word Cloud Generator")
        self.configure(bg=C["bg"])
        self.minsize(WINDOW_MIN_WIDTH, WINDOW_MIN_HEIGHT)
        self.update_idletasks()

        # Size to 80 % of the screen, centred, nudged up slightly for taskbars.
        sw, sh = self.winfo_screenwidth(), self.winfo_screenheight()
        w = max(WINDOW_MIN_WIDTH, int(sw * 0.80))
        h = max(WINDOW_MIN_HEIGHT, int(sh * 0.85))
        x = (sw - w) // 2
        y = max(0, (sh - h) // 2 - 30)
        self.geometry(f"{w}x{h}+{x}+{y}")

        # ── Application state ─────────────────────────────────────────────
        self.terms:          dict[str, int] = {}
        self.runner_ups:     dict[str, int] = {}
        self.user_excluded:  set[str]       = set()
        self._lock           = threading.Lock()
        self.detected_lang:  str  = "en"
        self.resume_text:    str  = ""
        self._file_path:     str  = ""   # full OS path used for file reading
        self.preview_image         = None  # holds ImageTk reference to prevent GC

        # Treeview sort state
        self._sort_col:     str  = "weight"
        self._sort_reverse: bool = True

        # Tk variables
        self.bg_var        = tk.StringVar(value="white")
        self.palette_var   = tk.StringVar(value="turbo")
        self.separator_var = tk.StringVar(value="_")
        self.separator_var.trace_add("write", self._on_separator_change)

        # Temp PNG written by Generate; cleaned up on close or regeneration.
        self._temp_banner_path: str | None = None

        # Widget references set in _build_ui, used for deferred show/pack.
        self._appearance_frame     = None
        self._action_buttons_frame = None

        self._build_ui()
        self._refresh_provider_label()
        self.protocol("WM_DELETE_WINDOW", self._on_close)

    # ── UI construction ───────────────────────────────────────────────────

    def _build_ui(self) -> None:
        """Construct and lay out all widgets.

        The bottom section (Appearance, action buttons, preview) is packed
        *before* the scrollable controls so tkinter reserves space for it
        even when the window is at minimum height.
        """
        style = ttk.Style(self)
        style.theme_use("clam")

        _f   = _FONT_FAMILY
        _fn  = (_f, 10)
        _fb  = (_f, 10, "bold")
        _fs  = (_f, 9)
        _fsi = (_f, 9, "italic")
        _fh  = (_f, 16, "bold")

        # ttk style definitions
        style.configure("Card.TLabelframe",
                        background=C["frame"], borderwidth=1, relief="solid")
        style.configure("Card.TLabelframe.Label",
                        background=C["frame"], foreground=C["text"], font=_fb)
        style.configure("TLabel",
                        background=C["bg"], foreground=C["text"], font=_fn)
        style.configure("Header.TLabel",
                        background=C["bg"], foreground=C["text"], font=_fh)
        style.configure("Sub.TLabel",
                        background=C["bg"], foreground=C["text_light"], font=_fs)
        style.configure("Status.TLabel",
                        background=C["bg"], foreground=C["text_light"], font=_fsi)
        style.configure("TRadiobutton",
                        background=C["frame"], foreground=C["text"], font=_fn)
        style.configure("TCombobox", font=_fn, padding=4)
        style.configure("TEntry", font=_fn,
                        fieldbackground=C["entry_bg"], foreground=C["text"])
        style.configure("Treeview",
                        background=C["tree_bg"], foreground=C["tree_fg"],
                        fieldbackground=C["tree_bg"], font=_fn, rowheight=24)
        style.configure("Treeview.Heading",
                        background=C["tree_head_bg"],
                        foreground=C["tree_head_fg"],
                        font=_fb, relief="flat", padding=4)
        style.map("Treeview.Heading",
                  background=[("active", C["border"])])
        style.map("Treeview",
                  background=[("selected", C["tree_sel"])],
                  foreground=[("selected", C["tree_fg"])])
        for btn_name, bg, hv, fg in [
            ("Primary.TButton", C["primary"], C["primary_hv"], "#ffffff"),
            ("Success.TButton", C["success"], C["success_hv"], "#ffffff"),
            ("Danger.TButton",  C["danger"],  C["danger_hv"],  "#ffffff"),
        ]:
            style.configure(btn_name, background=bg, foreground=fg,
                            font=(_f, 9, "bold"), padding=(14, 5),
                            relief="flat", borderwidth=0)
            style.map(btn_name,
                      background=[("active", hv), ("pressed", hv)],
                      foreground=[("active", fg), ("pressed", fg)])

        # ── Bottom panel (packed first to anchor at bottom) ───────────────
        preview_outer = tk.Frame(self, bg=C["bg"])
        preview_outer.pack(side="bottom", fill="x")

        # Appearance section — not packed yet; revealed in _populate_tree.
        af = ttk.LabelFrame(
            preview_outer, text="  Appearance  ", style="Card.TLabelframe",
        )
        self._appearance_frame = af  # pack deferred until first extraction
        inner = tk.Frame(af, bg=C["frame"])
        inner.pack(fill="x", padx=12, pady=6)

        tk.Label(inner, text="Background:", bg=C["frame"],
                 font=(_f, 10), fg=C["text"]).grid(
            row=0, column=0, sticky="w", padx=(0, 8))
        ttk.Radiobutton(inner, text="Light",
                        variable=self.bg_var, value="white").grid(
            row=0, column=1, padx=(0, 8))
        ttk.Radiobutton(inner, text="Dark",
                        variable=self.bg_var, value="black").grid(
            row=0, column=2, padx=(0, 24))
        tk.Label(inner, text="Palette:", bg=C["frame"],
                 font=(_f, 10), fg=C["text"]).grid(
            row=0, column=3, sticky="w", padx=(0, 8))

        palettes = [
            "turbo \u2014 Vibrant",   "gray \u2014 Mono",
            "ocean \u2014 Ocean",     "hot \u2014 Hot",
            "rainbow \u2014 Rainbow", "viridis \u2014 Viridis",
            "plasma \u2014 Plasma",   "inferno \u2014 Inferno",
        ]
        combo_pal = ttk.Combobox(inner, values=palettes,
                                 state="readonly", width=20)
        combo_pal.current(0)
        combo_pal.grid(row=0, column=4)
        combo_pal.bind(
            "<<ComboboxSelected>>",
            lambda e: self.palette_var.set(
                combo_pal.get().split(" \u2014 ")[0].strip()
            ),
        )

        tk.Label(inner, text="Separator:", bg=C["frame"],
                 font=(_f, 10), fg=C["text"]).grid(
            row=1, column=0, sticky="w", padx=(0, 8), pady=(8, 0))
        ttk.Radiobutton(inner, text="Underscore  (Data_Science)",
                        variable=self.separator_var, value="_").grid(
            row=1, column=1, columnspan=2, sticky="w",
            padx=(0, 8), pady=(8, 0))
        ttk.Radiobutton(inner, text="Space  (Data Science)",
                        variable=self.separator_var, value=" ").grid(
            row=1, column=3, sticky="w", padx=(0, 8), pady=(8, 0))
        tk.Label(
            inner,
            text="Affects both the treeview and the word cloud output.",
            bg=C["frame"], fg=C["text_light"], font=(_f, 8, "italic"),
        ).grid(row=2, column=0, columnspan=5, sticky="w",
               padx=(0, 8), pady=(2, 4))

        # Action buttons row
        bf = tk.Frame(preview_outer, bg=C["bg"])
        bf.pack(fill="x", padx=24, pady=(8, 4))
        self._action_buttons_frame = bf
        self._btn(bf, "Generate Word Cloud",
                  self._run_generate, C["primary"], width=22).pack(
            side="left", padx=(0, 8))
        self._btn(bf, "Save As \u2026",
                  self._save_as, C["primary"], width=14).pack(
            side="left", padx=(0, 8))
        self._btn(bf, "Export Terms (.txt)",
                  self._export_terms, C["primary"], width=18).pack(side="left")

        # Preview image
        pf = ttk.LabelFrame(preview_outer, text="  Preview  ",
                            style="Card.TLabelframe")
        pf.pack(fill="x", padx=24, pady=(0, 12))
        self.preview_label = tk.Label(
            pf, bg=C["frame"],
            text="Word cloud will appear here.",
            fg=C["text_light"], font=(_FONT_FAMILY, 10, "italic"),
            height=8,
        )
        self.preview_label.pack(fill="x", padx=8, pady=8)

        # ── Scrollable upper area ─────────────────────────────────────────
        scroll_outer = tk.Frame(self, bg=C["bg"])
        scroll_outer.pack(side="top", fill="both", expand=True)

        canvas = tk.Canvas(scroll_outer, bg=C["bg"], highlightthickness=0)
        vsb    = ttk.Scrollbar(scroll_outer, orient="vertical",
                               command=canvas.yview)
        self.scroll_frame = tk.Frame(canvas, bg=C["bg"])
        self.scroll_frame.bind(
            "<Configure>",
            lambda e: canvas.configure(scrollregion=canvas.bbox("all")),
        )
        win_id = canvas.create_window((0, 0), window=self.scroll_frame,
                                      anchor="nw")
        canvas.bind(
            "<Configure>",
            lambda e: canvas.itemconfig(win_id, width=e.width),
        )
        canvas.configure(yscrollcommand=vsb.set)
        canvas.pack(side="left", fill="both", expand=True)
        vsb.pack(side="right", fill="y")

        def _on_scroll(event: tk.Event) -> None:
            """Handle mousewheel scroll on macOS/Windows and X11 Linux."""
            if sys.platform == "darwin":
                canvas.yview_scroll(-event.delta, "units")
            else:
                canvas.yview_scroll(int(-1 * (event.delta / 120)), "units")

        canvas.bind_all("<MouseWheel>", _on_scroll)
        # X11 Linux fires Button-4/5 instead of MouseWheel.
        canvas.bind_all("<Button-4>", lambda e: canvas.yview_scroll(-1, "units"))
        canvas.bind_all("<Button-5>", lambda e: canvas.yview_scroll( 1, "units"))

        c = tk.Frame(self.scroll_frame, bg=C["bg"])
        c.pack(fill="both", expand=True, padx=24, pady=12)

        # Header row: title + Settings button
        hdr = tk.Frame(c, bg=C["bg"])
        hdr.pack(fill="x", pady=(0, 2))
        ttk.Label(hdr, text="LinkedIn Banner Word Cloud Generator",
                  style="Header.TLabel").pack(side="left", anchor="w")
        ttk.Button(hdr, text="\u2699  Settings",
                   style="Primary.TButton",
                   command=self._open_settings,
                   cursor="hand2").pack(side="right")

        self._provider_label_var = tk.StringVar()
        ttk.Label(c, textvariable=self._provider_label_var,
                  style="Sub.TLabel").pack(anchor="w", pady=(0, 10))

        # File selection row — buttons anchored right so long filenames
        # truncate in the label rather than pushing buttons off-screen.
        file_outer = tk.Frame(c, bg=C["bg"])
        file_outer.pack(fill="x", pady=(0, 4))

        btn_frame = tk.Frame(file_outer, bg=C["bg"])
        btn_frame.pack(side="right")
        self._btn(btn_frame, "Analyse Resume",
                  self._run_extraction, C["success"]).pack(
            side="right", padx=(8, 0))
        self._btn(btn_frame, "Browse \u2026",
                  self._browse_file, C["primary"]).pack(
            side="right", padx=(8, 0))

        self.file_display_var = tk.StringVar(value="No file selected")
        ttk.Label(file_outer, textvariable=self.file_display_var).pack(
            side="left", fill="x", expand=True, padx=(0, 8))

        self.status_var = tk.StringVar(value="Ready \u2014 select a resume file.")
        ttk.Label(c, textvariable=self.status_var,
                  style="Status.TLabel").pack(anchor="w", pady=(4, 8))

        # Extracted terms treeview
        lf = ttk.LabelFrame(c, text="  Extracted Terms  ",
                            style="Card.TLabelframe")
        lf.pack(fill="both", expand=True, pady=(0, 8))
        cols = ("term", "weight", "bar")
        self.tree = ttk.Treeview(lf, columns=cols, show="headings", height=12)
        self.tree.heading("term",   text="Term \u25bd",
                          command=lambda: self._sort_by("term"))
        self.tree.heading("weight", text="\u25bc Weight",
                          command=lambda: self._sort_by("weight"))
        self.tree.heading("bar",    text="Relative")
        self.tree.column("term",   width=280, anchor="w")
        self.tree.column("weight", width=80,  anchor="e")
        self.tree.column("bar",    width=200, anchor="w")
        tree_vsb = ttk.Scrollbar(lf, orient="vertical", command=self.tree.yview)
        self.tree.configure(yscrollcommand=tree_vsb.set)
        self.tree.pack(side="left", fill="both", expand=True,
                       padx=(8, 0), pady=8)
        tree_vsb.pack(side="right", fill="y", padx=(0, 8), pady=8)

        # Term editing controls
        ef = tk.Frame(c, bg=C["bg"])
        ef.pack(fill="x", pady=(0, 8))

        self._btn(ef, "\u2212 Remove Selected",
                  self._remove_selected, C["danger"]).grid(
            row=0, column=0, padx=(0, 16))
        tk.Label(ef, text="Add term:", bg=C["bg"],
                 font=(_f, 10), fg=C["text"]).grid(
            row=0, column=1, sticky="w", padx=(0, 6))
        self.add_entry = tk.Entry(
            ef, width=26,
            bg=C["entry_bg"], fg=C["text"],
            insertbackground=C["text"],
            relief="flat",
            highlightbackground=C["border"], highlightthickness=1,
            font=(_f, 10),
        )
        self.add_entry.grid(row=0, column=2, padx=(0, 6))
        self._btn(ef, "+ Add", self._add_term, C["success"]).grid(
            row=0, column=3)

        tk.Label(ef, text="Runner-ups:", bg=C["bg"],
                 font=(_f, 10), fg=C["text"]).grid(
            row=1, column=0, sticky="w", padx=(0, 6), pady=(8, 0))
        self.runner_combo = ttk.Combobox(ef, state="readonly", width=30)
        self.runner_combo.grid(row=1, column=1, columnspan=2,
                               sticky="w", padx=(0, 6), pady=(8, 0))
        self._btn(ef, "\u2191 Promote",
                  self._promote_runner, C["primary"]).grid(
            row=1, column=3, pady=(8, 0))

    @staticmethod
    def _btn(
        parent: tk.Widget,
        text: str,
        command,
        colour: str,
        width: int | None = None,
    ) -> ttk.Button:
        """Create and return a themed ``ttk.Button``.

        Maps *colour* (a hex string from ``C``) to the matching ttk style.
        """
        style_map = {
            C["primary"]: "Primary.TButton",
            C["danger"]:  "Danger.TButton",
            C["success"]: "Success.TButton",
        }
        kwargs: dict = dict(
            text=text,
            command=command,
            style=style_map.get(colour, "Primary.TButton"),
            cursor="hand2",
        )
        if width is not None:
            kwargs["width"] = width
        return ttk.Button(parent, **kwargs)

    # ── State / provider helpers ──────────────────────────────────────────

    def _refresh_provider_label(self) -> None:
        """Update the subtitle label to show the active provider name."""
        config = load_config()
        name   = config.get("active_provider", "")
        if name:
            self._provider_label_var.set(
                f"AI Provider: {name}  \u00b7  "
                "change anytime via \u2699 Settings"
            )
        else:
            self._provider_label_var.set(
                "\u26a0  No provider configured \u2014 "
                "click \u2699 Settings to add an API key."
            )

    def _open_settings(self) -> None:
        """Open the Settings dialog."""
        SettingsDialog(self)

    def _browse_file(self) -> None:
        """Open a file picker restricted to supported document types."""
        path = filedialog.askopenfilename(
            title="Select Resume / CV",
            initialdir=str(_downloads_folder()),
            filetypes=[
                ("Supported documents", "*.pdf *.docx *.txt"),
                ("PDF",        "*.pdf"),
                ("Word",       "*.docx"),
                ("Plain text", "*.txt"),
            ],
        )
        if not path:
            return
        if Path(path).suffix.lower() not in self._SUPPORTED_EXTENSIONS:
            messagebox.showerror(
                "Unsupported file type",
                f"'{Path(path).name}' is not a supported document type.\n\n"
                "Please select a PDF, Word (.docx), or plain text (.txt) file.",
            )
            return
        self._file_path = path
        self.file_display_var.set(_shorten_path(path))
        self._set_status(f"Selected: {path}")

    def _set_status(self, msg: str) -> None:
        """Update the status label and flush pending UI events."""
        self.status_var.set(msg)
        self.update_idletasks()

    def _on_separator_change(self, *_) -> None:
        """Re-format all stored terms to use the newly chosen separator.

        Called automatically by the ``trace_add`` on ``separator_var``
        whenever the user clicks a separator radio button. Splits each term
        on any existing separator (``_`` or space) and rejoins with the
        new one — no re-extraction needed.
        """
        sep = self.separator_var.get()
        with self._lock:
            if not self.terms and not self.runner_ups:
                return

            def _reformat(d: dict[str, int]) -> dict[str, int]:
                return {
                    sep.join(p for p in re.split(r"[ _]+", term) if p): w
                    for term, w in d.items()
                }

            self.terms      = _reformat(self.terms)
            self.runner_ups = _reformat(self.runner_ups)

        self._populate_tree()
        self._populate_runners()

    # ── Treeview management ───────────────────────────────────────────────

    def _populate_tree(self) -> None:
        """Rebuild the treeview from ``self.terms``.

        Also reveals the Appearance section the first time terms exist.
        """
        for row in self.tree.get_children():
            self.tree.delete(row)
        if not self.terms:
            return

        # Reveal the Appearance panel the first time we have terms.
        if (
            self._appearance_frame is not None
            and not self._appearance_frame.winfo_ismapped()
        ):
            self._appearance_frame.pack(
                fill="x", padx=24, pady=(8, 4),
                before=self._action_buttons_frame,
            )

        mx = max(self.terms.values())
        if self._sort_col == "term":
            ordered = sorted(
                self.terms.items(),
                key=lambda kv: kv[0].lower(),
                reverse=self._sort_reverse,
            )
        else:
            ordered = sorted(
                self.terms.items(),
                key=lambda kv: kv[1],
                reverse=self._sort_reverse,
            )

        for term, weight in ordered:
            bar = "\u2588" * (int((weight / mx) * 20) if mx else 0)
            self.tree.insert("", "end", values=(term, weight, bar))

    def _sort_by(self, col: str) -> None:
        """Toggle sort order on *col* and refresh the treeview."""
        if self._sort_col == col:
            self._sort_reverse = not self._sort_reverse
        else:
            self._sort_col     = col
            self._sort_reverse = col == "weight"  # weight defaults desc

        dn, up, no = "\u25bc", "\u25b2", "\u25bd"
        if self._sort_col == "term":
            self.tree.heading(
                "term",
                text=f"{dn if self._sort_reverse else up} Term",
            )
            self.tree.heading("weight", text=f"Weight {no}")
        else:
            self.tree.heading("term", text=f"Term {no}")
            self.tree.heading(
                "weight",
                text=f"{dn if self._sort_reverse else up} Weight",
            )
        self._populate_tree()

    # ── Extraction ────────────────────────────────────────────────────────

    def _run_extraction(self) -> None:
        """Validate inputs then kick off a background extraction thread."""
        if not self._file_path:
            messagebox.showwarning("No file", "Select a resume first.")
            return

        if Path(self._file_path).suffix.lower() not in self._SUPPORTED_EXTENSIONS:
            messagebox.showerror(
                "Unsupported file type",
                f"'{Path(self._file_path).name}' cannot be processed.\n\n"
                "Supported: PDF (.pdf), Word (.docx), plain text (.txt).",
            )
            return

        config        = load_config()
        provider_name = config.get("active_provider", "")
        api_key       = config.get("keys", {}).get(provider_name, "")

        if not provider_name or not api_key:
            messagebox.showwarning(
                "No provider",
                "No AI provider configured.\n\n"
                "Click \u2699 Settings to add an API key first.",
            )
            self._open_settings()
            return

        log.info("Extraction started: %s (provider: %s)",
                 self._file_path, provider_name)
        self._set_status("Reading file \u2026")

        def _worker() -> None:
            try:
                self.resume_text = extract_text_from_file(self._file_path)
                n = len(self.resume_text)
                self.after(0, self._set_status,
                           f"{n:,} chars read.  Validating \u2026")

                ok, reason = _validate_content(self.resume_text)
                if not ok:
                    # Ask the user before wasting an API call on bad content.
                    evt     = threading.Event()
                    proceed = {"v": False}

                    def _ask(r: str = reason) -> None:
                        proceed["v"] = messagebox.askyesno(
                            "Content Warning",
                            f"{r}\n\nSend to AI anyway?",
                        )
                        evt.set()

                    self.after(0, _ask)
                    evt.wait()
                    if not proceed["v"]:
                        self.after(0, self._set_status, "Cancelled.")
                        return

                self.after(0, self._set_status,
                           f"{n:,} chars.  Sending to {provider_name} \u2026")

                provider = get_provider(provider_name, api_key)
                terms, runner_ups, lang, model = extract_terms_ai(
                    self.resume_text,
                    provider=provider,
                    separator=self.separator_var.get(),
                    progress_callback=lambda m: self.after(
                        0, self._set_status, m
                    ),
                )
                with self._lock:
                    self.terms, self.runner_ups = terms, runner_ups
                    self.user_excluded = set()
                self.detected_lang = lang
                self.after(0, self._populate_tree)
                self.after(0, self._populate_runners)
                self.after(
                    0, self._set_status,
                    f"Done \u2014 {len(terms)} terms, "
                    f"{len(runner_ups)} runner-ups  (provider: {model})",
                )
            except Exception as exc:
                log.error("Extraction failed: %s", exc, exc_info=True)
                self.after(0, lambda: messagebox.showerror("Error", str(exc)))
                self.after(0, self._set_status, "Error.")

        threading.Thread(target=_worker, daemon=True).start()

    # ── Term editing ──────────────────────────────────────────────────────

    def _add_term(self) -> None:
        """Add one or more comma-separated terms from the entry field."""
        raw = self.add_entry.get().strip()
        if not raw:
            return
        sep = self.separator_var.get()
        with self._lock:
            for part in raw.split(","):
                cleaned = clean_term(part.strip(), separator=sep)
                if cleaned:
                    self.terms[cleaned] = MANUAL_ADD_WEIGHT
                    self.user_excluded.discard(cleaned)
        self.add_entry.delete(0, "end")
        self._populate_tree()

    def _remove_selected(self) -> None:
        """Remove highlighted rows and auto-backfill from runner-ups."""
        sel = self.tree.selection()
        if not sel:
            messagebox.showinfo("Nothing selected",
                                "Click a term in the table first.")
            return
        with self._lock:
            for iid in sel:
                term = self.tree.item(iid, "values")[0]
                self.terms.pop(term, None)
                self.user_excluded.add(term)

            # Refill gaps with the highest-weight eligible runner-up.
            while len(self.terms) < MAX_TERMS and self.runner_ups:
                best_term, best_w = None, -1
                for t, w in self.runner_ups.items():
                    if t not in self.user_excluded and w > best_w:
                        best_term, best_w = t, w
                if best_term is None:
                    break
                self.terms[best_term] = best_w
                del self.runner_ups[best_term]

        self._populate_tree()
        self._populate_runners()

    def _promote_runner(self) -> None:
        """Move the selected runner-up into the main terms list."""
        sel = self.runner_combo.get()
        if not sel:
            return
        match = re.match(r"^(.+?)\s*\((\d+)\)$", sel)
        if match:
            term, weight = match.group(1), int(match.group(2))
        else:
            term, weight = sel, MANUAL_ADD_WEIGHT
        with self._lock:
            self.terms[term] = weight
            self.runner_ups.pop(term, None)
            self.user_excluded.discard(term)
        self._populate_tree()
        self._populate_runners()

    def _populate_runners(self) -> None:
        """Refresh the runner-up combobox, hiding user-excluded terms."""
        items = [
            f"{t} ({w})"
            for t, w in sorted(
                self.runner_ups.items(), key=lambda kv: kv[1], reverse=True
            )
            if t not in self.user_excluded
        ]
        self.runner_combo["values"] = items
        if items:
            self.runner_combo.current(0)
        else:
            self.runner_combo.set("")

    # ── Generation + file I/O ─────────────────────────────────────────────

    def _run_generate(self) -> None:
        """Render the word cloud to a temp file in a background thread."""
        if not self.terms:
            messagebox.showwarning("No terms", "Analyse a resume first.")
            return
        log.info("Generation started: %d terms, bg=%s, palette=%s",
                 len(self.terms), self.bg_var.get(), self.palette_var.get())
        self._set_status("Generating \u2026")

        def _worker() -> None:
            try:
                # Write to a temp file so nothing lands in Downloads until
                # the user explicitly clicks Save As.
                fd, tmp = tempfile.mkstemp(
                    suffix=".png", prefix="linkedin_banner_"
                )
                os.close(fd)

                generate_wordcloud(
                    self.terms,
                    background_color=self.bg_var.get(),
                    colormap=self.palette_var.get(),
                    output_path=tmp,
                    lang=self.detected_lang,
                )

                # Delete the previous temp file once the new one is ready.
                old = self._temp_banner_path
                self._temp_banner_path = tmp
                if old and old != tmp:
                    try:
                        os.unlink(old)
                    except OSError:
                        pass

                mb = os.path.getsize(tmp) / (1024 * 1024)
                self.after(0, self._show_preview, tmp)
                self.after(
                    0, self._set_status,
                    f"Ready to save  ({mb:.2f} MB) \u2014 "
                    "click Save As \u2026 to keep it",
                )
            except Exception as exc:
                log.error("Generation failed: %s", exc, exc_info=True)
                self.after(0, lambda: messagebox.showerror("Error", str(exc)))
                self.after(0, self._set_status, "Error.")

        threading.Thread(target=_worker, daemon=True).start()

    def _show_preview(self, path: str) -> None:
        """Load *path* into the preview label, scaled to fit the panel."""
        img = Image.open(path)
        self.preview_label.update_idletasks()
        avail_w = max(600, self.preview_label.winfo_width() - 16)
        new_h   = int(img.height * avail_w / img.width)
        img = img.resize((avail_w, new_h), Image.LANCZOS)
        self.preview_image = ImageTk.PhotoImage(img)
        self.preview_label.configure(
            image=self.preview_image, text="", height=new_h,
        )

    def _save_as(self) -> None:
        """Copy the temp banner PNG to a user-chosen destination."""
        if not self._temp_banner_path or \
                not os.path.isfile(self._temp_banner_path):
            messagebox.showwarning("No banner", "Generate a word cloud first.")
            return
        dest = filedialog.asksaveasfilename(
            title="Save LinkedIn Banner",
            defaultextension=".png",
            filetypes=[("PNG image", "*.png")],
            initialfile="linkedin_banner.png",
            initialdir=str(_downloads_folder()),
        )
        if dest:
            shutil.copy2(self._temp_banner_path, dest)
            log.info("Saved to: %s", dest)
            self._set_status(f"Saved to: {dest}")

    def _on_close(self) -> None:
        """Delete the temp banner file (if any) then destroy the window."""
        if self._temp_banner_path:
            try:
                os.unlink(self._temp_banner_path)
            except OSError:
                pass
        self.destroy()

    def _export_terms(self) -> None:
        """Write main terms, runner-ups, and excluded terms to a .txt file."""
        if not self.terms:
            messagebox.showwarning("No terms", "Nothing to export.")
            return
        dest = filedialog.asksaveasfilename(
            defaultextension=".txt",
            filetypes=[("Text", "*.txt")],
            initialfile="extracted_terms.txt",
            initialdir=str(_downloads_folder()),
        )
        if not dest:
            return

        visible_runners = {
            t: w
            for t, w in self.runner_ups.items()
            if t not in self.user_excluded
        }
        with open(dest, "w", encoding="utf-8") as fh:
            fh.write(f"# Main Terms ({len(self.terms)})\n")
            for term, weight in sorted(
                self.terms.items(), key=lambda kv: kv[1], reverse=True
            ):
                fh.write(f"{term}\t{weight}\n")

            if visible_runners:
                fh.write(f"\n# Runner-Ups ({len(visible_runners)})\n")
                for term, weight in sorted(
                    visible_runners.items(), key=lambda kv: kv[1], reverse=True
                ):
                    fh.write(f"{term}\t{weight}\n")

            if self.user_excluded:
                fh.write(f"\n# Excluded ({len(self.user_excluded)})\n")
                for term in sorted(self.user_excluded):
                    fh.write(f"{term}\n")

        log.info("Exported terms to: %s", dest)
        self._set_status(f"Exported to: {dest}")


# ---------------------------------------------------------------------------
# Entry point
# ---------------------------------------------------------------------------
def main() -> None:
    """Launch the application."""
    log.info("Log file: %s", LOG_FILE)
    WordCloudApp().mainloop()


if __name__ == "__main__":
    main()
